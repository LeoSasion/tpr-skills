from __future__ import annotations

import argparse
import os
import re
import tempfile
from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from batch_common import (
    expected_card_map,
    read_manifest,
    require_delivery_choice,
    resolve_word_identifier_visibility,
    select_rows,
)
from verify_delivery import verify_card, verify_manifest_gate, verify_manifest_profile_binding

NUMBER_RE = re.compile(r"^(\d+)_.*_A4\.png$", re.IGNORECASE)
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
# Leave a small amount of vertical slack for WPS' required trailing layout
# paragraph.  Keeping the image just below the row height preserves the full
# A4 card while avoiding a spurious blank second page in WPS.
CARD_HEIGHT_MM = 146.25
ROW_HEIGHT_MM = 146.60
TABLE_WIDTH_TWIPS = 11906
COL_WIDTH_TWIPS = TABLE_WIDTH_TWIPS // 2
IDENTIFIER_CLEAR_TOP = 1450
IDENTIFIER_CLEAR_BOTTOM = 1504
IDENTIFIER_DARK_THRESHOLD = 180
IDENTIFIER_DARK_PIXEL_MINIMUM = 8


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Place verified A4 action cards into an A4 DOCX, four per page."
    )
    parser.add_argument("cards_dir", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--manifest", type=Path)
    parser.add_argument("--character-profile", type=Path)
    parser.add_argument("--start", type=int)
    parser.add_argument("--end", type=int)
    parser.add_argument("--id", dest="identifiers", action="append", default=[])
    parser.add_argument("--allow-unselected", action="store_true")
    parser.add_argument("--work-dir", type=Path)
    parser.add_argument(
        "--jpeg-quality",
        type=int,
        help="Embed print-optimized JPEG copies at this quality (recommended: 90).",
    )
    parser.add_argument(
        "--identifier-visibility",
        choices=["hidden", "shown"],
        help="Word-only identifier policy; defaults to the manifest value or hidden.",
    )
    return parser.parse_args()


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_geometry(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_TWIPS))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for _ in range(2):
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(COL_WIDTH_TWIPS))
        grid.append(grid_col)

    for row in table.rows:
        row.height = Mm(ROW_HEIGHT_MM)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Mm(105)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(COL_WIDTH_TWIPS))
            tc_w.set(qn("w:type"), "dxa")


def resolve_legacy_sources(cards_dir: Path, start: int, end: int) -> list[Path]:
    numbered: dict[int, list[Path]] = {}
    for path in cards_dir.glob("*.png"):
        match = NUMBER_RE.match(path.name)
        if match:
            numbered.setdefault(int(match.group(1)), []).append(path)
    sources = []
    for number in range(start, end + 1):
        matches = sorted(numbered.get(number, []))
        if len(matches) != 1:
            raise RuntimeError(f"Expected one card for {number:03d}, found {len(matches)}")
        sources.append(matches[0])
    if len(sources) % 4:
        raise RuntimeError("The selected range must contain a multiple of four cards")
    return sources


def validate_sources(sources: list[Path]):
    for source in sources:
        with Image.open(source) as image:
            if image.format != "PNG":
                raise RuntimeError(f"Unexpected format for {source.name}: {image.format}")
            if image.size != (1240, 1754):
                raise RuntimeError(f"Unexpected size for {source.name}: {image.size}")
            dpi = image.info.get("dpi")
            if not dpi or not all(149 <= float(value) <= 151 for value in dpi[:2]):
                raise RuntimeError(f"Unexpected or missing 150 DPI metadata for {source.name}: {dpi}")
            image.load()


def resolve_sources(args, cards_dir: Path):
    if args.manifest:
        fieldnames, all_rows = read_manifest(args.manifest)
        rows = select_rows(
            all_rows,
            start=args.start,
            end=args.end,
            identifiers=args.identifiers,
        )
        require_delivery_choice(rows, "word")
        expected = expected_card_map(cards_dir, rows)
        if not args.allow_unselected:
            actual = {path.name for path in cards_dir.glob("*.png")}
            wanted = set(expected)
            if actual != wanted:
                raise RuntimeError(
                    f"PNG set mismatch; missing={sorted(wanted - actual)}, "
                    f"extra={sorted(actual - wanted)}"
                )
        failures = verify_manifest_profile_binding(
            args.manifest,
            fieldnames,
            rows,
            args.character_profile,
        )
        for name, (path, row) in expected.items():
            if not path.is_file():
                failures.append(f"{name}: missing")
                continue
            failures.extend(f"{row['number']}: {error}" for error in verify_manifest_gate(row))
            failures.extend(
                f"{name}: {error}"
                for error in verify_card(path, row, allow_legacy_caption_check=False)
            )
        if failures:
            raise RuntimeError("Cards are not ready for Word:\n- " + "\n- ".join(failures))
        ordered = [(path, row) for path, row in expected.values()]
    else:
        if args.identifiers:
            raise RuntimeError("--id requires --manifest")
        if (args.start is None) != (args.end is None):
            raise RuntimeError("--start and --end must be supplied together")
        start = 1 if args.start is None else args.start
        end = 200 if args.end is None else args.end
        sources = resolve_legacy_sources(cards_dir, start, end)
        if not args.allow_unselected:
            all_pngs = {path.name for path in cards_dir.glob("*.png")}
            selected_pngs = {path.name for path in sources}
            if all_pngs != selected_pngs:
                raise RuntimeError(
                    f"PNG set mismatch; missing={sorted(selected_pngs - all_pngs)}, "
                    f"extra={sorted(all_pngs - selected_pngs)}"
                )
        validate_sources(sources)
        ordered = [
            (source, {"number": NUMBER_RE.match(source.name).group(1)}) for source in sources
        ]
    if len(ordered) % 4:
        raise RuntimeError("The selected range must contain a multiple of four cards")
    return ordered


def count_identifier_dark_pixels(image: Image.Image) -> int:
    grayscale = image.convert("L")
    region = grayscale.crop(
        (0, IDENTIFIER_CLEAR_TOP, grayscale.width, IDENTIFIER_CLEAR_BOTTOM)
    )
    pixels = (
        region.get_flattened_data()
        if hasattr(region, "get_flattened_data")
        else region.getdata()
    )
    return sum(pixel < IDENTIFIER_DARK_THRESHOLD for pixel in pixels)


def make_word_copy(
    source: Path,
    work_dir: Path,
    jpeg_quality: int | None,
    identifier_visibility: str,
) -> Path:
    if jpeg_quality is not None and not 70 <= jpeg_quality <= 100:
        raise ValueError("--jpeg-quality must be between 70 and 100")
    suffix = ".jpg" if jpeg_quality is not None else ".png"
    target = work_dir / f"{source.stem}{suffix}"
    if target.resolve() == source.resolve():
        raise ValueError(
            "Word work directory must differ from the source cards directory; "
            f"refusing to overwrite {source}"
        )
    alternate_media = [
        work_dir / f"{source.stem}{extension}"
        for extension in (".png", ".jpg", ".jpeg")
        if work_dir / f"{source.stem}{extension}" != target
        and (work_dir / f"{source.stem}{extension}").exists()
    ]
    if alternate_media:
        raise RuntimeError(
            "Word work directory contains stale alternate media for "
            f"{source.stem}: {[path.name for path in alternate_media]}; "
            "use a clean --work-dir"
        )
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=target.name + ".", suffix=".tmp", dir=work_dir
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        with Image.open(source) as opened:
            image = ImageOps.exif_transpose(opened).convert("RGB")
            if image.size != (1240, 1754):
                raise RuntimeError(f"Unexpected size for {source.name}: {image.size}")
            if identifier_visibility == "hidden":
                image.paste(
                    "white",
                    (0, IDENTIFIER_CLEAR_TOP, image.width, IDENTIFIER_CLEAR_BOTTOM),
                )
            if jpeg_quality is None:
                image.save(temporary, format="PNG", optimize=True, dpi=(150, 150))
                expected_format = "PNG"
            else:
                image.save(
                    temporary,
                    format="JPEG",
                    quality=jpeg_quality,
                    subsampling=0,
                    optimize=True,
                    dpi=(150, 150),
                )
                expected_format = "JPEG"

        with Image.open(temporary) as check:
            if check.format != expected_format or check.size != (1240, 1754):
                raise RuntimeError(
                    f"Unexpected Word media format/size for {source.name}: "
                    f"{check.format}/{check.size}"
                )
            check.load()
            dark_pixels = count_identifier_dark_pixels(check)
        if identifier_visibility == "hidden" and dark_pixels >= IDENTIFIER_DARK_PIXEL_MINIMUM:
            raise RuntimeError(f"Word media still shows an identifier: {source.name}")
        if identifier_visibility == "shown" and dark_pixels < IDENTIFIER_DARK_PIXEL_MINIMUM:
            raise RuntimeError(f"Word media identifier is not visible: {source.name}")
        os.chmod(temporary, 0o644)
        os.replace(temporary, target)
        return target
    finally:
        temporary.unlink(missing_ok=True)


def add_picture_border(inline_shape):
    sp_pr = inline_shape._inline.graphic.graphicData.pic.spPr
    line = OxmlElement("a:ln")
    line.set("w", "6350")  # 0.5 pt
    solid_fill = OxmlElement("a:solidFill")
    color = OxmlElement("a:srgbClr")
    color.set("val", "000000")
    solid_fill.append(color)
    line.append(solid_fill)
    dash = OxmlElement("a:prstDash")
    dash.set("val", "solid")
    line.append(dash)
    sp_pr.append(line)


def configure_page(document: Document):
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)


def add_page_break(document: Document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1)
    run = paragraph.add_run()
    run.font.size = Pt(1)
    run.add_break(WD_BREAK.PAGE)


def add_trailing_layout_paragraph(document: Document):
    """Keep WPS' required trailing paragraph inside the final A4 page."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1)
    run = paragraph.add_run()
    run.font.size = Pt(1)


def add_card_page(document: Document, sources: list[Path]):
    table = document.add_table(rows=2, cols=2)
    set_table_geometry(table)
    remove_table_borders(table)
    for cell, source in zip((cell for row in table.rows for cell in row.cells), sources):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Mm(0)
        paragraph.paragraph_format.right_indent = Mm(0)
        run = paragraph.add_run()
        inline_shape = run.add_picture(str(source), height=Mm(CARD_HEIGHT_MM))
        add_picture_border(inline_shape)


def main():
    args = parse_args()
    cards_dir = args.cards_dir.resolve()
    output = args.output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    resolved = resolve_sources(args, cards_dir)
    source_specs = [row for _, row in resolved]
    sources = [path for path, _ in resolved]
    identifier_visibility = resolve_word_identifier_visibility(
        source_specs, args.identifier_visibility
    )
    work_dir = (args.work_dir or output.parent / f"{output.stem}_assets").resolve()
    work_dir.mkdir(parents=True, exist_ok=True)
    converted = []
    for index, source in enumerate(sources, 1):
        converted.append(
            make_word_copy(
                source,
                work_dir,
                args.jpeg_quality,
                identifier_visibility,
            )
        )
        if index % 25 == 0 or index == len(sources):
            print(f"prepared Word media {index}/{len(sources)}", flush=True)
    sources = converted

    document = Document()
    configure_page(document)
    for offset in range(0, len(sources), 4):
        if offset:
            add_page_break(document)
        add_card_page(document, sources[offset:offset + 4])
        if (offset + 4) % 20 == 0 or offset + 4 == len(sources):
            print(f"embedded {offset + 4}/{len(sources)}", flush=True)
    add_trailing_layout_paragraph(document)

    core = document.core_properties
    first_identifier = source_specs[0]["number"]
    last_identifier = source_specs[-1]["number"]
    core.title = f"自适应角色TPR动作卡 {first_identifier}-{last_identifier} /img2word"
    visibility_label = "编号隐藏" if identifier_visibility == "hidden" else "编号显示"
    core.subject = f"A4纵向，每页4张，仅保留各图片细黑边框，{visibility_label}"
    core.author = ""
    core.last_modified_by = ""
    document.save(output)
    print(f"{output} ({identifier_visibility})", flush=True)


if __name__ == "__main__":
    main()
